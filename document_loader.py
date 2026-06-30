"""
Модуль для загрузки документов из папки knowledge_base.

Поддерживает различные форматы файлов:
- TXT (текстовые файлы)
- CSV (табличные данные)
- JSON (структурированные данные)
- PDF (текст из PDF файлов)
- DOCX (текст из документов Word)

Каждый документ получает метаданные о источнике (имя файла, путь).
"""

import csv
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Класс для загрузки документов из директории knowledge_base.
    
    Автоматически определяет формат файлов и извлекает текст.
    Каждый документ сохраняется с метаданными об источнике.
    """
    
    def __init__(self, knowledge_base_dir: str = "knowledge_base"):
        """
        Инициализация загрузчика документов.
        
        Args:
            knowledge_base_dir: Директория с файлами базы знаний
        """
        self.knowledge_base_dir = Path(knowledge_base_dir)
        
        logger.info(f"Инициализация загрузчика документов: {knowledge_base_dir}")
        
        # Создаем директорию, если она не существует
        if not self.knowledge_base_dir.exists():
            logger.warning(f"Директория базы знаний не найдена: {knowledge_base_dir}")
            logger.info("Создаю директорию knowledge_base/")
            self.knowledge_base_dir.mkdir(exist_ok=True)
    
    def load_all_documents(self) -> List[Tuple[str, str, Dict]]:
        """
        Загружает все документы из директории.
        
        Returns:
            Список кортежей (текст, источник, метаданные)
        """
        documents = []
        
        if not self.knowledge_base_dir.exists():
            logger.error(f"Директория базы знаний не существует: {self.knowledge_base_dir}")
            return documents
        
        # Рекурсивно обходим все файлы
        for file_path in self.knowledge_base_dir.rglob("*"):
            if file_path.is_file():
                try:
                    document_data = self._load_document(file_path)
                    if document_data:
                        text, source, metadata = document_data
                        if text.strip():  # Пропускаем пустые документы
                            documents.append((text, source, metadata))
                            logger.info(f"Загружен document: {source} ({len(text)} символов)")
                except Exception as e:
                    logger.error(f"Ошибка при загрузке файла {file_path.name}: {e}")
        
        logger.info(f"Всего загружено документов: {len(documents)}")
        return documents
    
    def _load_document(self, file_path: Path) -> Optional[Tuple[str, str, Dict]]:
        """
        Загружает один документ в зависимости от его формата.
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Кортеж (текст, источник, метаданные) или None
        """
        extension = file_path.suffix.lower()
        
        # Пропускаем системные файлы и уже обработанные документы
        if file_path.name.startswith('.') or file_path.name.startswith('_'):
            return None
        
        # Пропускаем файлы ChromaDB
        if 'chroma_db' in str(file_path).lower():
            return None
        
        try:
            if extension == '.txt':
                return self._load_txt(file_path)
            elif extension == '.csv':
                return self._load_csv(file_path)
            elif extension == '.json':
                return self._load_json(file_path)
            elif extension == '.pdf':
                return self._load_pdf(file_path)
            elif extension == '.docx':
                return self._load_docx(file_path)
            elif extension in ['.md', '.markdown']:
                return self._load_txt(file_path)  # Markdown как текст
            else:
                logger.debug(f"Неподдерживаемый формат файла: {file_path.name}")
                return None
                
        except Exception as e:
            logger.error(f"Ошибка при загрузке {file_path.name}: {e}")
            return None

    def _load_txt(self, file_path: Path) -> Tuple[str, str, Dict]:
        """Загружает текстовый файл."""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        source = str(file_path.relative_to(self.knowledge_base_dir))
        metadata = {
            "file_type": "txt",
            "file_size": file_path.stat().st_size,
            "file_path": str(file_path)
        }
        
        return text, source, metadata
    
    def _load_csv(self, file_path: Path) -> Optional[Tuple[str, str, Dict]]:
        """Загружает CSV файл, создавая отдельные документы для каждой строки."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                
                # Создаем документы из строк CSV
                for i, row in enumerate(reader, 1):
                    text_parts = []
                    for key, value in row.items():
                        if value and value.strip():
                            text_parts.append(f"{key}: {value}")
                    
                    if text_parts:
                        text = "\n".join(text_parts)
                        source = f"{file_path.name} (строка {i})"
                        metadata = {
                            "file_type": "csv",
                            "row_number": i,
                            "columns": list(row.keys()),
                            "file_path": str(file_path)
                        }
                        return text, source, metadata
                
        except Exception as e:
            logger.error(f"Ошибка при загрузке CSV {file_path.name}: {e}")
        
        return None
    
    def _load_json(self, file_path: Path) -> Optional[Tuple[str, str, Dict]]:
        """Загружает JSON файл."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Преобразуем JSON в читаемый текст
            text = self._json_to_text(data)
            
            if text.strip():
                source = str(file_path.relative_to(self.knowledge_base_dir))
                metadata = {
                    "file_type": "json",
                    "file_size": file_path.stat().st_size,
                    "file_path": str(file_path)
                }
                return text, source, metadata
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка парсинга JSON {file_path.name}: {e}")
        except Exception as e:
            logger.error(f"Ошибка при загрузке JSON {file_path.name}: {e}")
        
        return None
    
    def _json_to_text(self, data, prefix: str = "", depth: int = 0) -> str:
        """Рекурсивно преобразует JSON структуру в текст."""
        lines = []
        indent = "  " * depth
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{indent}{key}:")
                    lines.append(self._json_to_text(value, prefix, depth + 1))
                else:
                    lines.append(f"{indent}{key}: {value}")
        elif isinstance(data, list):
            for i, item in enumerate(data, 1):
                lines.append(f"{indent}{i}. {self._json_to_text(item, prefix, depth + 1)}")
        
        return "\n".join(lines)

    def _load_pdf(self, file_path: Path) -> Optional[Tuple[str, str, Dict]]:
        """Загружает PDF файл (требует библиотеку PyPDF2)."""
        try:
            import PyPDF2
            
            text_parts = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            if text_parts:
                text = "\n\n".join(text_parts)
                source = str(file_path.relative_to(self.knowledge_base_dir))
                metadata = {
                    "file_type": "pdf",
                    "pages": len(reader.pages),
                    "file_size": file_path.stat().st_size,
                    "file_path": str(file_path)
                }
                return text, source, metadata
        except ImportError:
            logger.warning(f"Библиотека PyPDF2 не установлена. Пропускаю PDF: {file_path.name}")
            logger.info("Установите PyPDF2: pip install PyPDF2")
        except Exception as e:
            logger.error(f"Ошибка при загрузке PDF {file_path.name}: {e}")
        
        return None
    
    def _load_docx(self, file_path: Path) -> Optional[Tuple[str, str, Dict]]:
        """Загружает DOCX файл (требует библиотеку python-docx)."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            if text_parts:
                text = "\n\n".join(text_parts)
                source = str(file_path.relative_to(self.knowledge_base_dir))
                metadata = {
                    "file_type": "docx",
                    "paragraphs": len(doc.paragraphs),
                    "file_size": file_path.stat().st_size,
                    "file_path": str(file_path)
                }
                return text, source, metadata
        except ImportError:
            logger.warning(f"Библиотека python-docx не установлена. Пропускаю DOCX: {file_path.name}")
            logger.info("Установите python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"Ошибка при загрузке DOCX {file_path.name}: {e}")
        
        return None
    
    def get_documents_by_type(self, file_type: str) -> List[Tuple[str, str, Dict]]:
        """
        Получает документы определенного типа.
        
        Args:
            file_type: Тип файла (txt, csv, json, pdf, docx)
            
        Returns:
            Список документов указанного типа
        """
        all_docs = self.load_all_documents()
        return [(text, source, meta) for text, source, meta in all_docs 
                if meta.get("file_type") == file_type]
